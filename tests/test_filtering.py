"""Tests for inclusion-based file filtering.

TDD: Written BEFORE the implementation.
Tests the INDEXED_EXTENSIONS allowlist, .gitignore parsing, and include CLI.
"""

from pathlib import Path

import pytest

import lss_store
from lss_store import _walk_text_files, _is_text_file, INDEXED_EXTENSIONS


# ── INDEXED_EXTENSIONS whitelist ─────────────────────────────────────────────


class TestIndexedExtensions:
    """Known text/code/doc extensions must be in the allowlist."""

    @pytest.mark.parametrize(
        "ext",
        [
            # Text / docs
            ".txt", ".md", ".rst", ".org", ".tex",
            # Code - Python
            ".py",
            # Code - JS/TS
            ".js", ".ts", ".jsx", ".tsx", ".mjs", ".cjs",
            # Code - systems
            ".go", ".rs", ".c", ".h", ".cpp", ".hpp", ".cc",
            # Code - JVM
            ".java", ".kt", ".scala",
            # Code - scripting
            ".rb", ".php", ".pl", ".lua", ".r",
            # Code - other
            ".swift", ".cs",
            # Shell
            ".sh", ".bash", ".zsh",
            # Web
            ".html", ".htm", ".css", ".scss",
            # Data
            ".json", ".jsonl", ".yaml", ".yml", ".toml", ".ini", ".csv", ".tsv",
            ".xml",
            # Config
            ".cfg", ".conf",
            # SQL
            ".sql", ".graphql",
            # Documents with extractors
            ".pdf", ".docx", ".xlsx", ".pptx", ".eml",
        ],
    )
    def test_known_extension_included(self, ext):
        assert ext in INDEXED_EXTENSIONS, f"{ext} should be in INDEXED_EXTENSIONS"

    @pytest.mark.parametrize(
        "ext",
        [
            # Binary formats
            ".png", ".jpg", ".gif", ".mp4", ".mp3", ".zip",
            # Compiled
            ".pyc", ".o", ".so", ".dll", ".wasm",
            # Unknown
            ".xyz", ".abc", ".foobar",
        ],
    )
    def test_binary_and_unknown_not_included(self, ext):
        assert ext not in INDEXED_EXTENSIONS, f"{ext} should NOT be in INDEXED_EXTENSIONS"


# ── File walking with inclusion filter ───────────────────────────────────────


class TestWalkInclusion:
    def test_known_extensions_indexed(self, tmp_path):
        """Files with known extensions are discovered."""
        (tmp_path / "readme.md").write_text("hello")
        (tmp_path / "app.py").write_text("print(1)")
        (tmp_path / "style.css").write_text("body{}")

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "readme.md" in names
        assert "app.py" in names
        assert "style.css" in names

    def test_unknown_extension_skipped(self, tmp_path):
        """Files with unknown extensions are NOT discovered."""
        (tmp_path / "data.xyz").write_text("some data")
        (tmp_path / "output.foobar").write_text("some output")
        (tmp_path / "readme.md").write_text("hello")

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "data.xyz" not in names
        assert "output.foobar" not in names
        assert "readme.md" in names  # this one should still be found

    def test_binary_extension_still_rejected(self, tmp_path):
        """Binary formats are fast-rejected (not even byte-checked)."""
        (tmp_path / "image.png").write_bytes(b"\x89PNG\r\n" + b"\x00" * 100)
        (tmp_path / "video.mp4").write_bytes(b"\x00" * 100)

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "image.png" not in names
        assert "video.mp4" not in names

    def test_docx_now_discoverable(self, tmp_path):
        """DOCX files (previously in BINARY_EXTENSIONS) are now discoverable."""
        from docx import Document

        docx_path = tmp_path / "report.docx"
        doc = Document()
        doc.add_paragraph("Important report content")
        doc.save(str(docx_path))

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "report.docx" in names

    def test_xlsx_now_discoverable(self, tmp_path):
        """XLSX files are now discoverable."""
        from openpyxl import Workbook

        xlsx_path = tmp_path / "data.xlsx"
        wb = Workbook()
        wb.active.append(["Name", "Value"])
        wb.save(str(xlsx_path))

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "data.xlsx" in names

    def test_pptx_now_discoverable(self, tmp_path):
        """PPTX files are now discoverable."""
        from pptx import Presentation

        pptx_path = tmp_path / "slides.pptx"
        Presentation().save(str(pptx_path))

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "slides.pptx" in names


# ── Known extensionless files ────────────────────────────────────────────────


class TestExtensionlessFiles:
    @pytest.mark.parametrize(
        "name",
        ["Makefile", "Dockerfile", "LICENSE", "README", "Procfile", "Vagrantfile",
         "Gemfile", "Rakefile", ".gitignore", ".dockerignore", ".editorconfig"],
    )
    def test_known_extensionless_included(self, tmp_path, name):
        """Well-known extensionless files should be discovered."""
        (tmp_path / name).write_text("content here")
        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert name in names, f"{name} should be discovered"

    def test_random_extensionless_skipped(self, tmp_path):
        """Random extensionless files with binary-ish content should be skipped."""
        # Create a file with no extension and null bytes (binary)
        (tmp_path / "mystery_binary").write_bytes(b"\x00\x01\x02\x03" * 100)
        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "mystery_binary" not in names


# ── .gitignore parsing ───────────────────────────────────────────────────────


class TestGitignoreParsing:
    def test_gitignore_excludes_files(self, tmp_path):
        """.gitignore patterns should exclude matching files."""
        (tmp_path / ".gitignore").write_text("*.log\nbuild/\n")
        (tmp_path / "app.py").write_text("print('hello')")
        (tmp_path / "server.log").write_text("log entry")
        build = tmp_path / "build"
        build.mkdir()
        (build / "output.js").write_text("compiled")

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "app.py" in names
        assert "server.log" not in names
        assert "output.js" not in names

    def test_nested_gitignore(self, tmp_path):
        """.gitignore in subdirectory only affects that subtree."""
        sub = tmp_path / "subdir"
        sub.mkdir()
        (sub / ".gitignore").write_text("*.tmp\n")
        (sub / "data.tmp").write_text("temporary")
        (sub / "code.py").write_text("real code")
        # Root-level .tmp should NOT be excluded by subdirectory .gitignore
        (tmp_path / "root.tmp").write_text("root temp")
        (tmp_path / "root.py").write_text("root code")

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "code.py" in names
        assert "data.tmp" not in names
        # .tmp extension is unknown, so root.tmp would be skipped by inclusion filter anyway
        assert "root.py" in names

    def test_gitignore_comments_and_blanks(self, tmp_path):
        """.gitignore comments (#) and blank lines are ignored."""
        (tmp_path / ".gitignore").write_text("# This is a comment\n\n*.log\n\n# Another comment\n")
        (tmp_path / "app.py").write_text("code")
        (tmp_path / "debug.log").write_text("log")

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "app.py" in names
        assert "debug.log" not in names

    def test_gitignore_directory_patterns(self, tmp_path):
        """.gitignore directory patterns (ending with /) exclude directories."""
        (tmp_path / ".gitignore").write_text("dist/\n")
        dist = tmp_path / "dist"
        dist.mkdir()
        (dist / "bundle.js").write_text("bundled")
        (tmp_path / "src.js").write_text("source")

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "src.js" in names
        assert "bundle.js" not in names


# ── User include extensions ──────────────────────────────────────────────────


class TestUserIncludeExtensions:
    def test_user_added_extension(self, tmp_path):
        """User-added extensions via config should be included."""
        import lss_config

        # Add .xyz to user include list
        cfg = lss_config.load_config()
        cfg["include_extensions"] = [".xyz"]
        lss_config.save_config(cfg)

        (tmp_path / "data.xyz").write_text("custom format data")
        (tmp_path / "readme.md").write_text("hello")

        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "data.xyz" in names
        assert "readme.md" in names

    def test_user_include_doesnt_affect_binary_rejection(self, tmp_path):
        """User include can't override binary extension rejection."""
        import lss_config

        # Even if user adds .png, it should still be rejected by BINARY_EXTENSIONS
        cfg = lss_config.load_config()
        cfg["include_extensions"] = [".png"]
        lss_config.save_config(cfg)

        (tmp_path / "image.png").write_bytes(b"\x89PNG\r\n" + b"\x00" * 100)
        files = list(_walk_text_files(tmp_path))
        names = {f.name for f in files}
        assert "image.png" not in names


# ── Integration: _is_text_file ───────────────────────────────────────────────


class TestIsTextFileUpdated:
    def test_docx_is_text(self, tmp_path):
        from docx import Document

        path = tmp_path / "test.docx"
        Document().save(str(path))
        assert _is_text_file(path) is True

    def test_xlsx_is_text(self, tmp_path):
        from openpyxl import Workbook

        path = tmp_path / "test.xlsx"
        Workbook().save(str(path))
        assert _is_text_file(path) is True

    def test_pptx_is_text(self, tmp_path):
        from pptx import Presentation

        path = tmp_path / "test.pptx"
        Presentation().save(str(path))
        assert _is_text_file(path) is True
